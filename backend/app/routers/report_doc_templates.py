"""
Report document templates — DOCX and Markdown with {{tag}} placeholders.

Supported tags
──────────────
Text tags (inline replacement):
  {{case.title}}              Case title
  {{case.id}}                 Case UUID
  {{case.status}}             Case status
  {{case.severity}}           Case severity (upper-case)
  {{case.tlp}}                TLP classification
  {{case.created_at}}         Creation date (YYYY-MM-DD HH:MM UTC)
  {{case.closed_at}}          Closure date  (or "N/A")
  {{case.description}}        Case description
  {{case.executive_summary}}  Executive summary
  {{case.quick_notes}}        Quick notes
  {{case.assigned_to}}        Assigned analyst(s)
  {{case.tags}}               Case tags (comma-separated)
  {{report.date}}             Report generation date (YYYY-MM-DD)
  {{report.author}}           Username of the analyst generating the report

Report content (analyst-authored, rendered as formatted DOCX paragraphs):
  {{report_analysis}}         Analyse Technique (box 1 of the Report tab)
  {{report_remediation}}      Remédiations      (box 2 of the Report tab)
  {{report_conclusion}}       Conclusion        (box 3 of the Report tab)
  {{report_content}}          All 3 boxes combined (backward compat alias)

Block tags (replaced with a table or image):
  {{ioc_table}}               Full IOC table
  {{asset_table}}             Full asset table
  {{evidence_table}}          Full evidence table
  {{timeline_table}}          Timeline (chronological)
  {{attack_graph}}            Attack-graph PNG image (DOCX) or placeholder (MD)
  {{mitre_matrix}}            MITRE ATT&CK coverage table (parents + selected sub-techniques)
  {{mitre_matrix_img}}        MITRE ATT&CK matrix as a visual PNG image (DOCX only)
"""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..database import get_db
from ..models.report_doc_template import ReportDocTemplate
from ..models.case import Case
from ..models.attack_graph import AttackGraph
from ..models.user import User
from ..core.deps import get_current_user
from ..config import settings

router = APIRouter(prefix="/report-doc-templates", tags=["report-doc-templates"])

# ── Storage directory ──────────────────────────────────────────────────────────

TEMPLATES_DIR: Path = settings.evidence_store_path.parent / "report_doc_templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

# ── Tag constants ──────────────────────────────────────────────────────────────

TAG_RE = re.compile(r'\{\{[\w.]+\}\}')

BLOCK_TAGS: set[str] = {
    "ioc_table", "asset_table", "evidence_table", "timeline_table",
    "attack_graph", "mitre_matrix", "mitre_matrix_img",
    # Report content tags (markdown → formatted DOCX paragraphs)
    "report_analysis", "report_remediation", "report_conclusion",
    "report_content",   # combined alias (backward compat)
}

ALL_TAGS: list[str] = [
    "case.title", "case.id", "case.status", "case.severity", "case.tlp",
    "case.created_at", "case.closed_at", "case.description",
    "case.executive_summary", "case.quick_notes", "case.assigned_to", "case.tags",
    "report.date", "report.author",
    # Report content (split)
    "report_analysis", "report_remediation", "report_conclusion",
    "report_content",  # combined backward compat
    # Annexes
    "ioc_table", "asset_table", "evidence_table", "timeline_table",
    "attack_graph", "mitre_matrix", "mitre_matrix_img",
]

# ── Pydantic schema ────────────────────────────────────────────────────────────

class ReportDocTemplateOut(BaseModel):
    id:            int
    name:          str
    description:   str
    format:        str
    file_size:     int
    tags_detected: list[str]
    created_at:    datetime
    created_by:    Optional[str]

    model_config = {"from_attributes": True}


# ── Tag detection ──────────────────────────────────────────────────────────────

def _detect_tags(text: str) -> list[str]:
    """Return sorted unique {{tag}} keys found in *text*."""
    return sorted({m[2:-2] for m in TAG_RE.findall(text)})


def _detect_tags_docx(file_bytes: bytes) -> list[str]:
    from docx import Document  # type: ignore
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append("".join(r.text for r in para.runs))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    parts.append("".join(r.text for r in para.runs))
    for section in doc.sections:
        for hdr_ftr in (section.header, section.footer):
            if hdr_ftr and not hdr_ftr.is_linked_to_previous:
                for para in hdr_ftr.paragraphs:
                    parts.append("".join(r.text for r in para.runs))
    return _detect_tags("\n".join(parts))


# ── Context builder ────────────────────────────────────────────────────────────

def _fmt_dt(dt: Optional[datetime]) -> str:
    return dt.strftime("%Y-%m-%d %H:%M UTC") if dt else "N/A"


def _build_context(case: Case, author: str) -> dict[str, str]:
    return {
        "case.title":             case.title or "",
        "case.id":                case.id or "",
        "case.status":            (case.status.value if case.status else "").replace("_", " ").title(),
        "case.severity":          (case.severity.value if case.severity else "").upper(),
        "case.tlp":               case.tlp or "",
        "case.created_at":        _fmt_dt(case.created_at),
        "case.closed_at":         _fmt_dt(case.closed_at),
        "case.description":       case.description or "",
        "case.executive_summary": case.executive_summary or "",
        "case.quick_notes":       case.quick_notes or "",
        "case.assigned_to":       case.assigned_to or "Unassigned",
        "case.tags":              case.tags or "",
        "report.date":            datetime.now(timezone.utc).strftime("%Y-%m-%d"),
        "report.author":          author,
    }


# ── Markdown helpers ───────────────────────────────────────────────────────────

def _md_ioc_table(case: Case) -> str:
    if not case.iocs:
        return "*No IOCs recorded.*"
    rows = [
        "| Type | Value | Confidence | TLP | Description |",
        "|------|-------|------------|-----|-------------|",
    ]
    for ioc in case.iocs:
        rows.append(
            f"| {ioc.type.value} | `{ioc.value}` | {ioc.confidence.value} "
            f"| {ioc.tlp} | {ioc.description or ''} |"
        )
    return "\n".join(rows)


def _md_asset_table(case: Case) -> str:
    if not case.assets:
        return "*No assets recorded.*"
    rows = [
        "| Name | Type | IP | Hostname | OS | Compromised |",
        "|------|------|----|----------|----|-------------|",
    ]
    for a in case.assets:
        rows.append(
            f"| {a.name} | {a.type.value} | {a.ip_address or ''} "
            f"| {a.hostname or ''} | {a.os or ''} | {'Yes' if a.compromised else 'No'} |"
        )
    return "\n".join(rows)


def _md_evidence_table(case: Case) -> str:
    if not case.evidences:
        return "*No evidence recorded.*"
    rows = [
        "| Name | File | Type | SHA-256 | Collected By |",
        "|------|------|------|---------|--------------|",
    ]
    for e in case.evidences:
        sha = (e.sha256_hash[:16] + "…") if e.sha256_hash else "N/A"
        rows.append(
            f"| {e.name} | {e.original_filename or ''} | {e.evidence_type.value} "
            f"| `{sha}` | {e.collected_by or ''} |"
        )
    return "\n".join(rows)


def _md_timeline_table(case: Case) -> str:
    if not case.timeline:
        return "*No timeline events recorded.*"
    rows = [
        "| Timestamp | Event | Actor | Source |",
        "|-----------|-------|-------|--------|",
    ]
    for ev in case.timeline:
        ts = ev.event_ts.strftime("%Y-%m-%d %H:%M UTC")
        rows.append(f"| {ts} | {ev.title} | {ev.actor or ''} | {ev.source or ''} |")
    return "\n".join(rows)


def _load_mitre_sub_map() -> dict[str, list[str]]:
    """Return {parent_id: [sub_id, ...]} from the local compact ATT&CK cache."""
    import json
    compact = settings.evidence_store_path.parent / "mitre" / "attack_enterprise_compact.json"
    if not compact.is_file():
        return {}
    sub_map: dict[str, list[str]] = {}
    data = json.loads(compact.read_text())
    for tactic in data.get("tactics", []):
        for tech in tactic.get("techniques", []):
            sub_map.setdefault(tech["id"], []).extend(
                s["id"] for s in tech.get("sub_techniques", [])
            )
    return sub_map


def _md_mitre_matrix(case: Case) -> str:
    """
    Render MITRE ATT&CK TTPs as a Markdown table.
    Rules:
    - Techniques with selected sub-techniques are "expanded": parent row + sub rows.
    - Techniques without selected sub-techniques show only the parent.
    - Sub-techniques are shown indented with ↳.
    """
    ttps = sorted(getattr(case, "ttps", []) or [], key=lambda t: (t.tactic or "", t.technique_id))
    if not ttps:
        return "*No MITRE ATT&CK techniques recorded.*"

    sub_map = _load_mitre_sub_map()
    selected_ids = {t.technique_id for t in ttps}

    rows = [
        "| Tactic | Technique ID | Technique Name |",
        "|--------|-------------|----------------|",
    ]
    emitted_parents: set[str] = set()

    for ttp in ttps:
        tid  = ttp.technique_id
        name = (ttp.technique_name or "").replace("|", "\\|")
        tact = (ttp.tactic_name or ttp.tactic or "Unknown").replace("|", "\\|")
        url  = f"https://attack.mitre.org/techniques/{tid.replace('.', '/')}/"
        link = f"[{tid}]({url})"
        is_sub = "." in tid

        if is_sub:
            # Sub-technique: ensure parent row appears first (once)
            parent_id = tid.rsplit(".", 1)[0]
            if parent_id not in emitted_parents:
                # Parent is not directly in the case — emit a placeholder row
                rows.append(f"| {tact} | {parent_id} | *(parent)* |")
                emitted_parents.add(parent_id)
            sub_link = f"[{tid}]({url})"
            rows.append(f"| | ↳ {sub_link} | {name} |")
        else:
            # Parent technique
            if tid in emitted_parents:
                continue  # already emitted as parent of a sub-tech
            emitted_parents.add(tid)
            has_selected_subs = any(s in selected_ids for s in sub_map.get(tid, []))
            rows.append(f"| {tact} | {link} | {name} |")
            if has_selected_subs:
                # Expanded: emit all selected sub-techniques inline
                for sub_id in sub_map.get(tid, []):
                    if sub_id in selected_ids:
                        sub_ttp = next((t for t in ttps if t.technique_id == sub_id), None)
                        sub_name = (sub_ttp.technique_name or "").replace("|", "\\|") if sub_ttp else ""
                        sub_url  = f"https://attack.mitre.org/techniques/{sub_id.replace('.', '/')}/"
                        rows.append(f"| | ↳ [{sub_id}]({sub_url}) | {sub_name} |")

    return "\n".join(rows)



# ── MITRE ATT&CK matrix image renderer ────────────────────────────────────────

# ATT&CK v19: defense-evasion split into stealth + defense-impairment
_TACTICS_V19 = [
    ("reconnaissance",       "Recon"),
    ("resource-development", "Rsrc Dev"),
    ("initial-access",       "Initial Access"),
    ("execution",            "Execution"),
    ("persistence",          "Persistence"),
    ("privilege-escalation", "Priv. Escalation"),
    ("stealth",              "Stealth"),
    ("defense-impairment",   "Def. Impairment"),
    ("credential-access",    "Cred. Access"),
    ("discovery",            "Discovery"),
    ("lateral-movement",     "Lateral Movement"),
    ("collection",           "Collection"),
    ("command-and-control",  "C2"),
    ("exfiltration",         "Exfiltration"),
    ("impact",               "Impact"),
]

_TACTIC_COLORS = {
    "reconnaissance":       "#a855f7",
    "resource-development": "#9333ea",
    "initial-access":       "#ef4444",
    "execution":            "#f97316",
    "persistence":          "#eab308",
    "privilege-escalation": "#f59e0b",
    "stealth":              "#84cc16",
    "defense-impairment":   "#d946ef",
    "credential-access":    "#22c55e",
    "discovery":            "#14b8a6",
    "lateral-movement":     "#06b6d4",
    "collection":           "#3b82f6",
    "command-and-control":  "#6366f1",
    "exfiltration":         "#8b5cf6",
    "impact":               "#f43f5e",
}


def _hex_to_rgb(h: str) -> tuple:
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) / 255.0 for i in (0, 2, 4))


def _render_mitre_matrix_png(case: Case) -> "Optional[tuple[bytes, float]]":
    """
    Render a clean, report-ready MITRE ATT&CK matrix image.

    Design:
    - White background (print-friendly), single horizontal row of tactic columns.
    - Column widths are auto-sized to fit their content (no text truncation).
    - Solid coloured tactic header bar with white bold text.
    - Each selected technique appears as a white card with a coloured
      left-accent stripe; sub-techniques are indented and slightly smaller.
    - Only the tactics that have ≥1 selected TTP are shown (simplified view).
    - Returns (png_bytes, embed_width_inches) so the DOCX can embed at the
      natural width (capped at 6.5" if the matrix is very wide).

    Note: TTPs saved with the legacy ``defense-evasion`` tactic slug (ATT&CK
    v18 and earlier) appear at the end under "Def. Evasion (legacy)".  To
    display them under the correct v19 Stealth / Defense Impairment columns
    the case TTPs must be deleted and re-added from the updated matrix.
    """
    try:
        import matplotlib                                   # type: ignore
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt                    # type: ignore
        from matplotlib.patches import Rectangle           # type: ignore

        ttps = sorted(
            getattr(case, "ttps", []) or [],
            key=lambda t: (t.tactic or "", t.technique_id or ""),
        )
        if not ttps:
            return None

        # ── Group TTPs by tactic, respecting ATT&CK v19 order ─────────────
        by_tactic: dict[str, list] = {}
        for ttp in ttps:
            by_tactic.setdefault(ttp.tactic or "", []).append(ttp)

        active = [(slug, label) for slug, label in _TACTICS_V19 if slug in by_tactic]
        known_slugs = {s for s, _ in _TACTICS_V19}
        for slug in by_tactic:          # legacy / unknown tactics at the end
            if slug not in known_slugs:
                if slug == "defense-evasion":
                    label = "Def. Evasion (legacy)"
                else:
                    label = slug.replace("-", " ").title()
                active.append((slug, label))
        if not active:
            return None

        # ── Layout constants (all in inches) ──────────────────────────────
        COL_GAP    = 0.08   # horizontal gap between adjacent columns
        HDR_H      = 0.38   # tactic header height
        CARD_H     = 0.36   # technique card height
        CARD_GAP   = 0.04   # vertical gap between cards
        MARGIN_X   = 0.15   # left/right page margin
        MARGIN_Y   = 0.15   # top/bottom page margin
        ACCENT_W   = 0.055  # width of the left colour stripe on cards
        INNER_X    = 0.075  # text left offset inside card (after accent stripe)
        SUB_INDENT = 0.10   # extra indent for sub-technique cards
        PAD_R      = 0.10   # right padding inside card before right edge

        # Colour palette
        BG_PAGE  = "#FFFFFF"
        BG_COL   = "#F8FAFC"   # slate-50 — column background
        BG_CARD  = "#FFFFFF"   # technique card
        BG_SUB   = "#F0FDF4"   # sub-technique card — light green tint
        TXT_ID   = "#1E293B"   # slate-800 — technique ID
        TXT_NAME = "#475569"   # slate-600 — technique name
        TXT_SUB  = "#15803D"   # green-700 — sub-technique name
        BORDER   = "#E2E8F0"   # slate-200

        # ── Auto-size each column to fit its content (no truncation) ──────
        # Approximate character widths at the rendered font sizes
        CHAR_MONO = 0.056   # monospace 5.6 pt  (technique IDs)
        CHAR_REG  = 0.046   # regular   5.0 pt  (technique names)
        CHAR_HDR  = 0.058   # bold      6.0 pt  (header labels)
        MIN_COL_W = 0.72    # absolute minimum column width

        col_widths: list[float] = []
        for slug, label in active:
            techs = by_tactic.get(slug, [])
            needed = 0.0
            for ttp in techs:
                tid    = ttp.technique_id or ""
                name   = (ttp.technique_name or "").strip()
                is_sub = "." in tid
                indent = SUB_INDENT if is_sub else 0.0
                text_x = ACCENT_W + INNER_X + indent
                w_id   = text_x + len(tid)  * CHAR_MONO + PAD_R
                w_name = text_x + len(name) * CHAR_REG  + PAD_R
                needed = max(needed, w_id, w_name)
            # Also ensure header text fits
            w_hdr = len(label) * CHAR_HDR + 0.22
            col_widths.append(max(MIN_COL_W, needed, w_hdr))

        N = len(active)
        fig_w = 2 * MARGIN_X + sum(col_widths) + COL_GAP * max(N - 1, 0)

        # Cap figure width at 14" to stay reasonable
        MAX_FIG_W = 14.0
        if fig_w > MAX_FIG_W:
            scale = MAX_FIG_W / fig_w
            col_widths = [w * scale for w in col_widths]
            fig_w = MAX_FIG_W

        # Height: tallest column determines figure height
        max_techs = max((len(by_tactic[slug]) for slug, _ in active), default=0)
        body_h = HDR_H + CARD_GAP + max_techs * (CARD_H + CARD_GAP)
        fig_h  = MARGIN_Y + body_h + MARGIN_Y

        DPI = 200

        fig, ax = plt.subplots(figsize=(fig_w, fig_h))
        fig.patch.set_facecolor(BG_PAGE)
        ax.set_facecolor(BG_PAGE)
        ax.set_xlim(0, fig_w)
        ax.set_ylim(fig_h, 0)   # y=0 at top, increases downward
        ax.axis("off")

        cursor_x = MARGIN_X
        for col_i, (slug, short_label) in enumerate(active):
            cx0   = cursor_x
            col_w = col_widths[col_i]
            color = _TACTIC_COLORS.get(slug, "#6b7280")
            r, g, b = _hex_to_rgb(color)

            # ── Column background ──────────────────────────────────────────
            ax.add_patch(Rectangle(
                (cx0, MARGIN_Y), col_w, body_h,
                facecolor=BG_COL, edgecolor=BORDER, linewidth=0.5, zorder=1,
            ))

            # ── Tactic header bar ──────────────────────────────────────────
            ax.add_patch(Rectangle(
                (cx0, MARGIN_Y), col_w, HDR_H,
                facecolor=(r, g, b), edgecolor="none", zorder=2,
            ))
            # Header text — break into two lines if the label is long
            words = short_label.split()
            if len(words) > 2 and len(short_label) > 13:
                mid   = len(words) // 2
                line1 = " ".join(words[:mid])
                line2 = " ".join(words[mid:])
                ax.text(cx0 + col_w / 2, MARGIN_Y + HDR_H / 2 - 0.055, line1,
                        color="white", fontsize=5.8, fontweight="bold",
                        ha="center", va="center", zorder=3)
                ax.text(cx0 + col_w / 2, MARGIN_Y + HDR_H / 2 + 0.055, line2,
                        color="white", fontsize=5.8, fontweight="bold",
                        ha="center", va="center", zorder=3)
            else:
                ax.text(cx0 + col_w / 2, MARGIN_Y + HDR_H / 2, short_label,
                        color="white", fontsize=6.0, fontweight="bold",
                        ha="center", va="center", zorder=3)

            # ── Technique cards ────────────────────────────────────────────
            techs = sorted(by_tactic.get(slug, []), key=lambda t: t.technique_id or "")
            for row_i, ttp in enumerate(techs):
                ty0    = MARGIN_Y + HDR_H + CARD_GAP + row_i * (CARD_H + CARD_GAP)
                is_sub = "." in (ttp.technique_id or "")
                indent = SUB_INDENT if is_sub else 0.0
                card_x = cx0 + indent
                card_w = col_w - indent

                # Card background
                ax.add_patch(Rectangle(
                    (card_x, ty0), card_w, CARD_H,
                    facecolor=BG_SUB if is_sub else BG_CARD,
                    edgecolor=BORDER, linewidth=0.4, zorder=2,
                ))
                # Left colour accent stripe
                ax.add_patch(Rectangle(
                    (card_x, ty0), ACCENT_W, CARD_H,
                    facecolor=(r, g, b, 0.55 if is_sub else 0.85),
                    edgecolor="none", zorder=3,
                ))

                text_x = card_x + ACCENT_W + INNER_X
                tid    = ttp.technique_id or ""
                name   = (ttp.technique_name or "").strip()

                # Technique ID (monospace)
                ax.text(text_x, ty0 + 0.085, tid,
                        color=TXT_SUB if is_sub else TXT_ID,
                        fontsize=5.2 if is_sub else 5.6,
                        fontweight="bold", ha="left", va="top",
                        zorder=4, fontfamily="monospace")

                # Technique name — no truncation (column is sized to fit)
                ax.text(text_x, ty0 + 0.205, name,
                        color=TXT_SUB if is_sub else TXT_NAME,
                        fontsize=4.8 if is_sub else 5.0,
                        ha="left", va="top", zorder=4)

            cursor_x += col_w + COL_GAP

        buf = io.BytesIO()
        plt.savefig(buf, format="png", bbox_inches="tight",
                    facecolor=BG_PAGE, dpi=DPI)
        plt.close(fig)
        buf.seek(0)

        # Embed width: natural size, capped at 6.5" (Word page text area)
        embed_w = min(fig_w, 6.5)
        return buf.read(), embed_w

    except Exception:
        return None


def _render_markdown(template_text: str, case: Case, ctx: dict[str, str]) -> str:
    text = template_text
    for tag, value in ctx.items():
        text = text.replace(f"{{{{{tag}}}}}", str(value))
    text = text.replace("{{ioc_table}}", _md_ioc_table(case))
    text = text.replace("{{asset_table}}", _md_asset_table(case))
    text = text.replace("{{evidence_table}}", _md_evidence_table(case))
    text = text.replace("{{timeline_table}}", _md_timeline_table(case))
    text = text.replace("{{attack_graph}}", "_[Attach the attack graph image — export it from the Attack Graph tab]_")
    text = text.replace("{{mitre_matrix}}", _md_mitre_matrix(case))
    text = text.replace("{{mitre_matrix_img}}", "_[MITRE ATT&CK matrix image — available in DOCX export only]_")
    # Split report content tags (fixed 3-section backward compat)
    text = text.replace("{{report_analysis}}",    (case.report_analysis    or "").strip() or "_[Aucune analyse rédigée.]_")
    text = text.replace("{{report_remediation}}", (case.report_remediation or "").strip() or "_[Aucune remédiation rédigée.]_")
    text = text.replace("{{report_conclusion}}",  (case.report_conclusion  or "").strip() or "_[Aucune conclusion rédigée.]_")
    # Dynamic per-section tags from report_sections_data
    import json as _json
    try:
        sections_data: dict = _json.loads(getattr(case, "report_sections_data", None) or "{}")
    except Exception:
        sections_data = {}
    for slug, content in sections_data.items():
        text = text.replace(f"{{{{{slug}}}}}", (content or "").strip() or f"_[Section «{slug}» non rédigée.]_")
    # Combined backward compat
    text = text.replace("{{report_content}}", case.report or "_[Aucun contenu de rapport rédigé.]_")
    return text


# ── DOCX helpers ───────────────────────────────────────────────────────────────

def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _replace_text_in_para(para, ctx: dict[str, str]) -> bool:
    """Merge all runs into one and replace {{tags}}. Returns True if changed."""
    from docx.oxml.ns import qn  # type: ignore
    full = _para_text(para)
    new_text = full
    for tag, value in ctx.items():
        new_text = new_text.replace(f"{{{{{tag}}}}}", str(value))
    if new_text == full:
        return False
    # Remove all existing <w:r> elements
    for r_elem in list(para._p.findall(qn("w:r"))):
        para._p.remove(r_elem)
    if new_text:
        para.add_run(new_text)
    return True


def _set_cell_bg(cell, color_hex: str) -> None:
    """Apply background shading to a table cell (no leading #)."""
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#").upper())
    tcPr.append(shd)


def _build_word_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    header_color: str = "1E3A5F",
) -> "Table":  # type: ignore[name-defined]
    """
    Create a fully styled Word table:
    - Extends ~0.5 cm beyond the page text-area on both sides
    - Solid borders (outer 1pt, inner 0.5pt) in neutral grey
    - Coloured header row with white bold text (colour per table type)
    - Alternating light-grey row shading for readability
    - Font size 9pt header / 8pt body
    """
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore

    EMU_PER_TWIP = 635        # 914 400 EMU/in ÷ 1 440 twips/in
    EXT_TWIPS    = 284         # ~0.5 cm extension into each margin
    BORDER_COLOR = "9CA3AF"   # tailwind gray-400 — neutral on any background

    n_cols = len(headers)
    n_rows = len(rows)
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)

    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # ── Width: extend beyond the text area ───────────────────────────────────
    col_w_twips: Optional[int] = None
    try:
        section  = doc.sections[0]
        text_w   = int((section.page_width - section.left_margin - section.right_margin)
                       / EMU_PER_TWIP)
        table_w  = text_w + 2 * EXT_TWIPS

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(table_w))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)

        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), str(-EXT_TWIPS))
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)

        col_w_twips = table_w // n_cols
    except Exception:
        # Fallback: 100 % of text area, no extension
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblPr.append(tblW)

    # ── Borders ───────────────────────────────────────────────────────────────
    tblBorders = OxmlElement("w:tblBorders")
    for side, sz in [
        ("top", "12"), ("left", "12"), ("bottom", "12"), ("right", "12"),
        ("insideH", "6"), ("insideV", "4"),
    ]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)       # half-points: 12 = 1.5 pt, 6 = 0.75 pt
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), BORDER_COLOR)
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # ── tblGrid: even column widths ───────────────────────────────────────────
    if col_w_twips:
        tblGrid = OxmlElement("w:tblGrid")
        for _ in headers:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(col_w_twips))
            tblGrid.append(gc)
        tblPr.addnext(tblGrid)

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        _set_cell_bg(cell, header_color)
        # Set column width on header cell too
        if col_w_twips:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_w_twips))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # ── Data rows ─────────────────────────────────────────────────────────────
    ALT_BG = "F3F4F6"   # tailwind gray-100 — subtle alternating shade
    for r_i, row_data in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row_data):
            cell = cells[c_i]
            if r_i % 2 == 1:
                _set_cell_bg(cell, ALT_BG)
            if col_w_twips:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(col_w_twips))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)
            para = cell.paragraphs[0]
            run = para.add_run(str(val) if val is not None else "")
            run.font.size = Pt(8)

    return table


def _build_mitre_word_table(doc, case: Case) -> "Table":  # type: ignore[name-defined]
    """
    Build a MITRE ATT&CK coverage Word table for the case.

    Layout: Tactic | Technique ID | Technique Name
    Rules:
    - Parent techniques are shown as regular rows.
    - If a parent has selected sub-techniques → it is "expanded":
      sub-technique rows appear immediately below, indented with ↳.
    - If a parent has no selected sub-techniques → only parent row shown.
    - Free-standing sub-techniques (selected without parent) appear indented
      after an auto-inserted parent placeholder row.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEADER_COLOR = "0F2942"   # Deep navy — neutral DFIR feel
    SUB_BG       = "F0FDF4"   # Very light green — sub-technique rows
    ALT_BG       = "F3F4F6"   # Light gray — alternating parent rows

    sub_map = _load_mitre_sub_map()
    ttps = sorted(getattr(case, "ttps", []) or [], key=lambda t: (t.tactic or "", t.technique_id))

    if not ttps:
        # Return an empty placeholder paragraph cast as a table-like object
        # by abusing _build_word_table with empty rows
        return _build_word_table(doc, ["Tactic", "Technique ID", "Technique Name"], [], HEADER_COLOR)

    selected_ids = {t.technique_id for t in ttps}
    ttp_by_id    = {t.technique_id: t for t in ttps}

    # Build ordered display rows: (tactic_name, tech_id, tech_name, is_sub)
    display: list[tuple[str, str, str, bool]] = []
    emitted: set[str] = set()

    for ttp in ttps:
        tid  = ttp.technique_id
        tact = ttp.tactic_name or ttp.tactic or "Unknown"
        name = ttp.technique_name or ""
        is_sub = "." in tid

        if is_sub:
            parent_id = tid.rsplit(".", 1)[0]
            if parent_id not in emitted:
                # Emit parent placeholder if not already in the case
                if parent_id in ttp_by_id:
                    parent = ttp_by_id[parent_id]
                    display.append((parent.tactic_name or parent.tactic or tact, parent_id, parent.technique_name or "", False))
                else:
                    display.append((tact, parent_id, "", False))
                emitted.add(parent_id)
            display.append((tact, tid, name, True))
            emitted.add(tid)
        else:
            if tid in emitted:
                continue
            emitted.add(tid)
            display.append((tact, tid, name, False))
            # Immediately after parent: emit any selected sub-techniques
            for sub_id in sub_map.get(tid, []):
                if sub_id in selected_ids and sub_id not in emitted:
                    sub_ttp = ttp_by_id.get(sub_id)
                    sub_name = sub_ttp.technique_name or "" if sub_ttp else ""
                    display.append((tact, sub_id, sub_name, True))
                    emitted.add(sub_id)

    # ── Build the Word table ──────────────────────────────────────────────────
    n_rows = len(display)
    table  = doc.add_table(rows=1 + n_rows, cols=3)
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Width: full page text area
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    # Borders
    BORDER = "9CA3AF"
    tblBorders = OxmlElement("w:tblBorders")
    for side, sz in [("top","12"),("left","12"),("bottom","12"),("right","12"),("insideH","6"),("insideV","4")]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0");    b.set(qn("w:color"), BORDER)
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Column widths (approx): Tactic 25%, ID 20%, Name 55%
    COL_PCTS = [1250, 1000, 2750]  # out of 5000 twips
    tblGrid = OxmlElement("w:tblGrid")
    for w in COL_PCTS:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); tblGrid.append(gc)
    tblPr.addnext(tblGrid)

    def _set_col_w(cell, w: int) -> None:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)

    # Header row
    hdr = table.rows[0].cells
    for i, (label, w) in enumerate(zip(["Tactic", "Technique ID", "Technique Name"], COL_PCTS)):
        _set_cell_bg(hdr[i], HEADER_COLOR); _set_col_w(hdr[i], w)
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(9)

    # Data rows
    prev_tactic = None
    for r_i, (tact, tid, name, is_sub) in enumerate(display):
        cells = table.rows[r_i + 1].cells
        for c_i, w in enumerate(COL_PCTS):
            _set_col_w(cells[c_i], w)

        if is_sub:
            # Sub-technique row: light green background, indented ID
            for c in cells: _set_cell_bg(c, SUB_BG.lstrip("#"))
            cells[0].paragraphs[0].add_run("").font.size = Pt(8)   # blank tactic cell
            id_run  = cells[1].paragraphs[0].add_run(f"  ↳ {tid}")
            name_run = cells[2].paragraphs[0].add_run(name)
            for r in (id_run, name_run):
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x16, 0x6A, 0x34)  # forest green
        else:
            # Parent technique row: alternate shading
            if r_i % 2 == 1:
                for c in cells: _set_cell_bg(c, ALT_BG.lstrip("#"))
            tact_text = tact if tact != prev_tactic else ""  # only show tactic on first row of group
            prev_tactic = tact
            for c_i, (text, w) in enumerate(zip([tact_text, tid, name], COL_PCTS)):
                run = cells[c_i].paragraphs[0].add_run(text)
                run.font.size = Pt(8)
                if c_i == 1:  # technique ID — slightly bold
                    run.bold = True

    return table


def _render_attack_graph_png(nodes: list, edges: list) -> Optional[bytes]:
    """
    Render the attack graph using the actual ReactFlow node positions and types,
    reproducing the visual style of the AttackGraph tab.

    Nodes are stored as ReactFlow Node objects:
      { id, type, position: {x, y}, data: {label, subLabel, nodeKind, notes,
        compromised, iocType, ...}, style: {width: 200} }
    Edges: { id, source, target, ... }
    """
    try:
        import matplotlib                           # type: ignore
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt            # type: ignore
        import matplotlib.patches as mpatches      # type: ignore
        from matplotlib.patches import FancyBboxPatch  # type: ignore

        if not nodes:
            return None

        # ── Constants matching the frontend ───────────────────────────────────
        NODE_W  = 200    # matches NODE_WIDTH in AttackGraphNodes.tsx
        NODE_H  = 72     # estimated — label (18) + subLabel (14) + padding (40)
        BG      = "#0B121F"

        # Border + fill colours per nodeKind
        THEME = {
            "timeline":           {"border": "#9FEF00", "fill": "#0e1f05"},
            "asset":              {"border": "#3b82f6", "fill": "#050f20"},
            "asset_compromised":  {"border": "#ef4444", "fill": "#1a0505"},
            "attacker":           {"border": "#ef4444", "fill": "#1a0505"},
            "free":               {"border": "#4b5563", "fill": "#111827"},
        }
        IOC_COLORS = {
            "ip":          "#ef4444",
            "domain":      "#f97316",
            "url":         "#fb923c",
            "hash_md5":    "#a855f7",
            "hash_sha1":   "#a855f7",
            "hash_sha256": "#a855f7",
            "email":       "#3b82f6",
            "filename":    "#eab308",
            "registry":    "#ec4899",
            "certificate": "#ec4899",
            "email_subject": "#60a5fa",
        }

        node_map: dict[str, dict] = {str(n["id"]): n for n in nodes}

        # ── Compute bounding box from real positions ────────────────────────
        xs = [n.get("position", {}).get("x", 0) for n in nodes]
        ys = [n.get("position", {}).get("y", 0) for n in nodes]
        pad = 80
        min_x, max_x = min(xs) - pad, max(xs) + NODE_W + pad
        min_y, max_y = min(ys) - pad, max(ys) + NODE_H + pad
        data_w = max(max_x - min_x, 1)
        data_h = max(max_y - min_y, 1)

        # Scale to a reasonable figure size (1 in ≈ 120 data units, cap at 20 in)
        SCALE = 120
        fig_w = max(10.0, min(20.0, data_w / SCALE))
        fig_h = max(6.0,  min(14.0, data_h / SCALE))

        fig, ax = plt.subplots(figsize=(fig_w, fig_h))
        fig.patch.set_facecolor(BG)
        ax.set_facecolor(BG)
        # ReactFlow: Y increases downward → invert matplotlib Y axis
        ax.set_xlim(min_x, max_x)
        ax.set_ylim(max_y, min_y)
        ax.axis("off")

        # ── Draw edges ─────────────────────────────────────────────────────
        for edge in (edges or []):
            src_node = node_map.get(str(edge.get("source", "")))
            tgt_node = node_map.get(str(edge.get("target", "")))
            if not src_node or not tgt_node:
                continue
            sp = src_node.get("position", {"x": 0, "y": 0})
            tp = tgt_node.get("position", {"x": 0, "y": 0})
            x1 = sp["x"] + NODE_W / 2
            y1 = sp["y"] + NODE_H       # bottom-centre of source
            x2 = tp["x"] + NODE_W / 2
            y2 = tp["y"]                # top-centre of target
            ax.annotate(
                "", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(
                    arrowstyle="-|>",
                    color="#4b5563",
                    lw=1.2,
                    connectionstyle="arc3,rad=0.04",
                ),
                zorder=1,
            )

        # ── Draw nodes ─────────────────────────────────────────────────────
        for node in nodes:
            pos  = node.get("position", {"x": 0, "y": 0})
            data = node.get("data", {})
            ntype = str(node.get("type", "free"))
            x, y = pos["x"], pos["y"]

            # Pick theme
            if ntype == "attacker":
                th = THEME["attacker"]
            elif ntype == "timeline":
                th = THEME["timeline"]
            elif ntype == "asset":
                th = THEME["asset_compromised"] if data.get("compromised") else THEME["asset"]
            elif ntype == "ioc":
                col = IOC_COLORS.get(str(data.get("iocType", "")), "#a855f7")
                th = {"border": col, "fill": "#0d0516"}
            else:  # free / unknown
                th = THEME["free"]

            # Rounded rectangle
            rect = FancyBboxPatch(
                (x + 1, y + 1), NODE_W - 2, NODE_H - 2,
                boxstyle="round,pad=0,rounding_size=6",
                facecolor=th["fill"],
                edgecolor=th["border"],
                linewidth=1.8,
                zorder=2,
            )
            ax.add_patch(rect)

            # Attacker node: pill shape label centred
            label     = str(data.get("label", node.get("id", "")))
            sub_label = str(data.get("subLabel", ""))
            notes     = str(data.get("notes", ""))

            # Truncate long strings
            if len(label) > 32:
                label = label[:30] + "…"
            if len(sub_label) > 36:
                sub_label = sub_label[:34] + "…"
            if len(notes) > 52:
                notes = notes[:50] + "…"

            text_x = x + 10
            if sub_label:
                ax.text(text_x, y + 12, sub_label,
                        color=th["border"], fontsize=6.0, va="top", ha="left",
                        fontstyle="italic", zorder=3)
                ax.text(text_x, y + 26, label,
                        color="white", fontsize=7.5, fontweight="bold",
                        va="top", ha="left", zorder=3)
                if notes:
                    ax.text(text_x, y + 47, notes,
                            color="#8b949e", fontsize=5.5, va="top", ha="left", zorder=3)
            else:
                ax.text(text_x, y + NODE_H / 2, label,
                        color="white", fontsize=7.5, fontweight="bold",
                        va="center", ha="left", zorder=3)
                if notes:
                    ax.text(text_x, y + NODE_H / 2 + 12, notes,
                            color="#8b949e", fontsize=5.5, va="top", ha="left", zorder=3)

        buf = io.BytesIO()
        plt.savefig(buf, format="png", bbox_inches="tight",
                    facecolor=BG, dpi=150)
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    except Exception:
        return None


def _resolve_image_url(url: str) -> "Optional[Path]":
    """
    Map a note-image URL (or vault view URL) to an absolute local path.

    Handles:
      /note-images/{case_id}/{filename}   → note_images dir on disk
      /api/v1/vaults/{id}/view            → vault file path from DB (not resolved here —
                                            caller should pass the path directly if needed)
    Returns None if the file cannot be resolved or does not exist.
    """
    base = settings.evidence_store_path.parent

    # /note-images/{case_id}/{filename}
    if url.startswith("/note-images/"):
        rel  = url[len("/note-images/"):]          # e.g. "abc123/image.png"
        path = base / "note_images" / rel
        return path if path.exists() else None

    # Absolute filesystem path (edge case — tiptap sometimes emits data URIs
    # or local paths; ignore those silently)
    return None


def _md_to_docx_paragraphs(doc, placeholder_para, md_text: str) -> None:
    """
    Replace *placeholder_para* with markdown rendered as proper DOCX content.

    Supported syntax:
      • # / ## / ### headings          → Heading 1 / 2 / 3 styles
      • ``` fenced code blocks         → Courier New paragraph
      • **bold**, *italic*, `code`     → character formatting
      • - / * bullet lists             → List Bullet style
      • 1. numbered lists              → List Number style
      • --- horizontal rules           → blank separator paragraph
      • ![alt](url)  standalone line   → embedded picture paragraph
      • ![alt](url)  inline in text    → embedded picture + surrounding text
        (falls back to [alt] if the image cannot be resolved)
    """
    import re as _re
    from docx.shared import Inches  # type: ignore

    # Regex that matches a full markdown image token
    IMG_RE = _re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
    # A line that is *only* an image (optional surrounding whitespace)
    ONLY_IMG_RE = _re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")

    ref    = placeholder_para._p
    parent = ref.getparent()

    def _insert_before(para):
        parent.remove(para._p)
        ref.addprevious(para._p)

    def _make_para(style: str = "Normal"):
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
        _insert_before(p)
        return p

    def _embed_image(para, url: str, max_w: float = 5.5) -> bool:
        """Try to embed an image from *url* into *para*. Returns True on success."""
        path = _resolve_image_url(url)
        if not path:
            return False
        try:
            para.add_run().add_picture(str(path), width=Inches(max_w))
            return True
        except Exception:
            return False

    def _add_inline(para, text: str) -> None:
        """
        Add runs to *para* parsing **bold**, *italic*, `code`, and ![img](url).
        Images found inline are embedded as picture runs; if embedding fails
        the alt text is inserted as a plain run.

        IMG_RE.split gives: [text_before, alt1, url1, text_between, alt2, url2, …, text_after]
        We iterate in steps of 3: plain text chunk, then optional (alt, url) image pair.
        """
        segments = IMG_RE.split(text)
        # segments = [text_before, alt, url, text_between, alt, url, …, text_after]
        idx = 0
        while idx < len(segments):
            chunk = segments[idx]
            # Inline formatting in the text chunk
            fmt_parts = _re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", chunk)
            for fp in fmt_parts:
                if not fp:
                    continue
                if fp.startswith("**") and fp.endswith("**"):
                    r = para.add_run(fp[2:-2]); r.bold = True
                elif fp.startswith("*") and fp.endswith("*"):
                    r = para.add_run(fp[1:-1]); r.italic = True
                elif fp.startswith("`") and fp.endswith("`"):
                    r = para.add_run(fp[1:-1]); r.font.name = "Courier New"
                else:
                    if fp:
                        para.add_run(fp)

            # If there is a following (alt, url) image pair, embed it
            if idx + 2 < len(segments):
                alt = segments[idx + 1]
                url = segments[idx + 2]
                ok  = _embed_image(para, url)
                if not ok:
                    para.add_run(f"[{alt}]")  # fallback: alt text
                idx += 3
            else:
                idx += 1

    lines      = (md_text or "").split("\n")
    in_code    = False
    code_lines: list[str] = []

    for line in lines:
        # ── Fenced code block ──────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                p = _make_para()
                r = p.add_run("\n".join(code_lines))
                r.font.name = "Courier New"
                in_code = False
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        # ── Standalone image line → own paragraph ─────────────────────────
        m = ONLY_IMG_RE.match(line)
        if m:
            alt, url = m.group(1), m.group(2)
            p = _make_para()
            ok = _embed_image(p, url)
            if not ok:
                para = p
                para.add_run(f"[Image: {alt or url}]").italic = True
            continue

        # ── Headings ───────────────────────────────────────────────────────
        if line.startswith("### "):
            p = _make_para("Heading 3"); _add_inline(p, line[4:])
        elif line.startswith("## "):
            p = _make_para("Heading 2"); _add_inline(p, line[3:])
        elif line.startswith("# "):
            p = _make_para("Heading 1"); _add_inline(p, line[2:])
        # ── Horizontal rule ────────────────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            _make_para()
        # ── Bullet list ────────────────────────────────────────────────────
        elif _re.match(r"^[-*+] ", line):
            p = _make_para("List Bullet"); _add_inline(p, line[2:])
        elif _re.match(r"^\d+\. ", line):
            p = _make_para("List Number"); _add_inline(p, _re.sub(r"^\d+\. ", "", line, count=1))
        # ── Non-empty line (may contain inline images) ─────────────────────
        elif line.strip():
            p = _make_para(); _add_inline(p, line)
        # ── Empty line → implicit paragraph break (no-op) ──────────────────

    # Remove the original placeholder paragraph
    parent.remove(ref)


def _render_docx(template_path: str, case: Case, ctx: dict[str, str],
                 attack_graph_png: Optional[bytes]) -> bytes:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    doc = Document(template_path)

    # ── Table data ─────────────────────────────────────────────────────────────
    ioc_headers = ["Type", "Value", "Confidence", "TLP", "Description"]
    ioc_rows = (
        [[i.type.value, i.value, i.confidence.value, i.tlp, i.description or ""]
         for i in case.iocs]
        if case.iocs else [["No IOCs recorded", "", "", "", ""]]
    )

    asset_headers = ["Name", "Type", "IP", "Hostname", "OS", "Compromised"]
    asset_rows = (
        [[a.name, a.type.value, a.ip_address or "", a.hostname or "",
          a.os or "", "Yes" if a.compromised else "No"]
         for a in case.assets]
        if case.assets else [["No assets recorded", "", "", "", "", ""]]
    )

    evidence_headers = ["Name", "File", "Type", "SHA-256", "Collected By"]
    evidence_rows = (
        [[e.name, e.original_filename or "", e.evidence_type.value,
          (e.sha256_hash[:16] + "…") if e.sha256_hash else "N/A",
          e.collected_by or ""]
         for e in case.evidences]
        if case.evidences else [["No evidence recorded", "", "", "", ""]]
    )

    timeline_headers = ["Timestamp", "Event", "Actor", "Source"]
    timeline_rows = (
        [[ev.event_ts.strftime("%Y-%m-%d %H:%M UTC"), ev.title,
          ev.actor or "", ev.source or ""]
         for ev in case.timeline]
        if case.timeline else [["No timeline events recorded", "", "", ""]]
    )

    # ── First pass: body paragraphs ────────────────────────────────────────────
    # Collect block-tag paragraphs; replace text in all others.
    block_paras: list[tuple] = []
    for para in list(doc.paragraphs):
        full = _para_text(para)
        block_found = next((bt for bt in BLOCK_TAGS if f"{{{{{bt}}}}}" in full), None)
        if block_found:
            block_paras.append((para, block_found))
        else:
            _replace_text_in_para(para, ctx)

    # ── Table cells (text only, no block tags) ─────────────────────────────────
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_text_in_para(para, ctx)

    # ── Headers / footers (text only) ─────────────────────────────────────────
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf and not hf.is_linked_to_previous:
                for para in hf.paragraphs:
                    _replace_text_in_para(para, ctx)
                for tbl in hf.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _replace_text_in_para(para, ctx)

    # ── Second pass: replace block-tag paragraphs ──────────────────────────────
    for para, block_tag in block_paras:

        if block_tag == "report_analysis":
            _md_to_docx_paragraphs(doc, para, (case.report_analysis or "").strip() or "_[Aucune analyse rédigée.]_")

        elif block_tag == "report_remediation":
            _md_to_docx_paragraphs(doc, para, (case.report_remediation or "").strip() or "_[Aucune remédiation rédigée.]_")

        elif block_tag == "report_conclusion":
            _md_to_docx_paragraphs(doc, para, (case.report_conclusion or "").strip() or "_[Aucune conclusion rédigée.]_")

        elif block_tag.startswith("report_") is False and block_tag not in (
            "ioc_table", "asset_table", "evidence_table", "timeline_table",
            "attack_graph", "mitre_matrix", "mitre_matrix_img", "report_content",
        ):
            # Dynamic per-section tag — look up in report_sections_data
            import json as _json
            try:
                sd: dict = _json.loads(getattr(case, "report_sections_data", None) or "{}")
            except Exception:
                sd = {}
            content = sd.get(block_tag, "").strip()
            _md_to_docx_paragraphs(doc, para, content or f"_[Section «{block_tag}» non rédigée.]_")

        elif block_tag == "report_content":
            # Combined backward compat alias
            _md_to_docx_paragraphs(doc, para, case.report or "")

        elif block_tag in ("attack_graph", "mitre_matrix_img"):
            # Image blocks — clear runs and embed a PNG picture
            for r_elem in list(para._p.findall(qn("w:r"))):
                para._p.remove(r_elem)

            if block_tag == "attack_graph":
                png        = attack_graph_png
                embed_w    = 6.0
                placeholder = "[Attack graph not available — no data recorded]"
            else:  # mitre_matrix_img
                result     = _render_mitre_matrix_png(case)
                png        = result[0] if result else None
                embed_w    = result[1] if result else 6.5
                placeholder = "[MITRE ATT&CK matrix — no techniques recorded]"

            if png:
                para.add_run().add_picture(io.BytesIO(png), width=Inches(embed_w))
            else:
                para.add_run(placeholder)

        else:
            # Table blocks
            if block_tag == "ioc_table":
                tbl = _build_word_table(doc, ioc_headers, ioc_rows, "7F1D1D")
            elif block_tag == "asset_table":
                tbl = _build_word_table(doc, asset_headers, asset_rows, "1E3A5F")
            elif block_tag == "evidence_table":
                tbl = _build_word_table(doc, evidence_headers, evidence_rows, "1E293B")
            elif block_tag == "mitre_matrix":
                tbl = _build_mitre_word_table(doc, case)
            else:  # timeline_table
                tbl = _build_word_table(doc, timeline_headers, timeline_rows, "14532D")

            # Move the new table (currently at end of body) to replace the paragraph.
            para._p.addnext(tbl._tbl)
            para._p.getparent().remove(para._p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Routes ─────────────────────────────────────────────────────────────────────

@router.get("/", response_model=list[ReportDocTemplateOut])
def list_templates(db: Session = Depends(get_db)):
    return (
        db.query(ReportDocTemplate)
        .order_by(ReportDocTemplate.created_at.desc())
        .all()
    )


@router.get("/tags", response_model=list[str])
def list_available_tags():
    """Return all supported {{tags}}."""
    return ALL_TAGS


@router.post("/upload", response_model=ReportDocTemplateOut)
async def upload_template(
    name:         str        = Form(...),
    description:  str        = Form(""),
    file:         UploadFile = File(...),
    db:           Session    = Depends(get_db),
    current_user: User       = Depends(get_current_user),
):
    fname = file.filename or ""
    if fname.endswith(".docx"):
        fmt = "docx"
    elif fname.endswith(".md"):
        fmt = "markdown"
    else:
        raise HTTPException(400, "Only .docx and .md files are accepted")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(400, "Uploaded file is empty")

    # Detect tags
    if fmt == "docx":
        tags = _detect_tags_docx(file_bytes)
    else:
        tags = _detect_tags(file_bytes.decode("utf-8", errors="replace"))

    # Persist to disk
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-.]", "_", fname)
    dest = TEMPLATES_DIR / f"{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}_{safe_name}"
    dest.write_bytes(file_bytes)

    tpl = ReportDocTemplate(
        name=name,
        description=description,
        format=fmt,
        file_path=str(dest),
        file_size=len(file_bytes),
        tags_detected=tags,
        created_by=current_user.username,
    )
    db.add(tpl)
    db.commit()
    db.refresh(tpl)
    return tpl


@router.delete("/{template_id}")
def delete_template(
    template_id:  int,
    db:           Session = Depends(get_db),
    current_user: User    = Depends(get_current_user),
):
    tpl = db.query(ReportDocTemplate).filter(ReportDocTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "Template not found")
    try:
        Path(tpl.file_path).unlink(missing_ok=True)
    except Exception:
        pass
    db.delete(tpl)
    db.commit()
    return {"deleted": template_id}


@router.post("/{template_id}/generate/{case_id}")
def generate_report(
    template_id:  int,
    case_id:      str,
    db:           Session = Depends(get_db),
    current_user: User    = Depends(get_current_user),
):
    tpl = db.query(ReportDocTemplate).filter(ReportDocTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(404, "Template not found")
    if not Path(tpl.file_path).exists():
        raise HTTPException(410, "Template file missing on disk")

    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(404, "Case not found")

    ctx = _build_context(case, current_user.username)
    safe_title = re.sub(r"[^\w\-]", "_", case.title or "report")

    if tpl.format == "markdown":
        template_text = Path(tpl.file_path).read_text(encoding="utf-8")
        content = _render_markdown(template_text, case, ctx)
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}_report.md"'},
        )

    # DOCX
    ag = db.query(AttackGraph).filter(AttackGraph.case_id == case_id).first()
    png_bytes: Optional[bytes] = None
    if ag and ag.nodes:
        png_bytes = _render_attack_graph_png(ag.nodes, ag.edges or [])

    docx_bytes = _render_docx(tpl.file_path, case, ctx, png_bytes)
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{safe_title}_report.docx"'},
    )
